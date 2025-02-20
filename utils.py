import os
import logging
from pathlib import Path
from typing import Optional, List, Dict, Any, Union
import json
import docx
from pypdf import PdfReader
from werkzeug.utils import secure_filename
from weaviate.connect import ConnectionParams
from langchain_ollama import OllamaEmbeddings
from weaviate.auth import AuthApiKey
import weaviate
from llama_index.core import VectorStoreIndex
from llama_index.core.node_parser import SentenceSplitter
from llama_index.vector_stores.weaviate import WeaviateVectorStore
from llama_index.embeddings.langchain import LangchainEmbedding
import redis
from statistics import mean
from collections import defaultdict
from concurrent.futures import ThreadPoolExecutor
from llama_index.readers.file.docs import DocxReader, PDFReader
from llama_index.readers.file.flat import FlatReader
from llama_index.readers.file.tabular import PandasCSVReader, PandasExcelReader
from llama_index.readers.file.slides import PptxReader
import datetime
from config import (
    OLLAMA_BASE_URL,
    OLLAMA_MODEL,
    UPLOAD_FOLDER,
    ALLOWED_EXTENSIONS,
    REDIS_URL
)
from llama_index.core.settings import Settings

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class FileHandler:
    CHUNK_SIZE = 4096  # Optimal chunk size for file reading
    
    @staticmethod
    def allowed_file(filename: str, allowed_extensions: set) -> bool:
        """Check if the file extension is allowed."""
        return '.' in filename and filename.rsplit('.', 1)[1].lower() in allowed_extensions

    @staticmethod
    def secure_file_save(file: Any, upload_folder: str) -> Optional[str]:
        """Securely save a file and return the file path."""
        try:
            filename = secure_filename(file.filename)
            logger.info(f"Saving file: {filename} to {upload_folder}")
            file_path = os.path.join(upload_folder, filename)
            file.save(file_path)
            logger.info(f"File saved successfully: {file_path}")
            return file_path
        except Exception as e:
            logger.error(f"Error saving file: {str(e)}")
            return None

    @staticmethod
    def read_file_content(file_path: str) -> Optional[str]:
        """Optimized file reading with proper error handling and memory management."""
        logger.info(f"Reading file content from: {file_path}")
        try:
            file_extension = Path(file_path).suffix.lower()[1:]
            
            if file_extension == 'txt':
                # Use generator for large files
                def read_in_chunks(file_obj):
                    while True:
                        data = file_obj.read(FileHandler.CHUNK_SIZE)
                        if not data:
                            break
                        yield data
                
                with open(file_path, 'r', encoding='utf-8') as file:
                    return ''.join(chunk for chunk in read_in_chunks(file))
            
            elif file_extension == 'json':
                with open(file_path, 'r') as file:
                    data = json.load(file)
                    # Handle both single objects and arrays
                    if isinstance(data, list):
                        return '\n'.join(json.dumps(item) for item in data)
                    return json.dumps(data)
            
            elif file_extension == 'pdf':
                text_chunks = []
                with open(file_path, 'rb') as file:
                    pdf_reader = PdfReader(file)
                    # Process pages in parallel for large PDFs
                    with ThreadPoolExecutor() as executor:
                        text_chunks = list(executor.map(
                            lambda page: page.extract_text(), 
                            pdf_reader.pages
                        ))
                return '\n'.join(chunk for chunk in text_chunks if chunk)
            
            elif file_extension == 'docx':
                doc = docx.Document(file_path)
                # Include tables and other elements
                text_parts = []
                for para in doc.paragraphs:
                    text_parts.append(para.text)
                for table in doc.tables:
                    for row in table.rows:
                        text_parts.append(' | '.join(cell.text for cell in row.cells))
                return '\n'.join(text_parts)
            
            else:
                raise ValueError(f"Unsupported file type: {file_extension}")
                
        except Exception as e:
            logger.error(f"Error reading file {file_path}: {str(e)}")
            return None

class WeaviateHelper:
    _client = None

    @classmethod
    def get_client(cls):
        if cls._client is None:
            cls._client = weaviate.connect_to_local(
                host="localhost",
                port=8080,
                grpc_port=50051
            )
            logger.info("Connected to local Weaviate instance.")
        return cls._client

    @classmethod
    def close_client(cls):
        if cls._client is not None:
            cls._client.close()
            cls._client = None
            logger.info("Weaviate client connection closed.")

    @staticmethod
    def check_connection() -> bool:
        """Check if Weaviate connection is working."""
        client = WeaviateHelper.get_client()
        try:
            if not client.is_ready():
                logger.warning("Weaviate client is not ready, attempting to reconnect...")
                # WeaviateHelper.close_client()  # Close the existing client
                WeaviateHelper.get_client()  # Reconnect
            logger.info("Successfully connected to Weaviate!")
            return True
        except Exception as e:
            logger.error(f"Failed to connect to Weaviate: {str(e)}")
            return False

    @staticmethod
    # def store_document_chunk(chunk: str, embedding: List[float], source: str, chunk_index: int) -> None:
    #     """Store a document chunk in Weaviate."""
    #     weaviate_client = WeaviateHelper.get_client()
    #     weaviate_client.data.create(
    #         data={
    #             "content": chunk,
    #             "source": source,
    #             "chunk_index": chunk_index,
    #             "status": "processed"
    #         },
    #         class_name="Document",
    #         vector=embedding
    #     )

    @staticmethod
    def initialize_client(
        rest_url: str,
        grpc_url: str,
        client_name: str,
        api_key: str,
        http_port: int = 8080,
        grpc_port: int = 50051,
        secure: bool = True
    ) -> Any:
        """Initialize and return a Weaviate client with standard configuration."""
        try:
            connection_params = ConnectionParams.from_params(
                http_host=rest_url,
                http_port=http_port,
                http_secure=secure,
                grpc_host=grpc_url,
                grpc_port=grpc_port,
                grpc_secure=secure
            )

            auth_config = AuthApiKey(api_key=api_key)

            client = weaviate.WeaviateClient(
                connection_params=connection_params,
                auth_client_secret=auth_config,
                additional_headers={
                    "X-Weaviate-Client-Name": client_name
                }
            )
            
            # Test connection
            if WeaviateHelper.check_connection():
                return client
            return None
            
        except Exception as e:
            logger.error(f"Error initializing Weaviate client: {str(e)}")
            return None

class TimeTracker:
    @staticmethod
    def track_time(func):
        """Decorator to track execution time of functions."""
        def wrapper(*args, **kwargs):
            import time
            start_time = time.time()
            result = func(*args, **kwargs)
            end_time = time.time()
            logger.info(f"{func.__name__} took {end_time - start_time:.2f} seconds")
            return result
        return wrapper 

class JSONQueryHelper:
    @staticmethod
    def extract_json_fields(text: str) -> Optional[Dict]:
        """Extract JSON data from text."""
        try:
            # Find JSON-like content in text
            start = text.find('{')
            end = text.rfind('}') + 1
            if start >= 0 and end > start:
                json_str = text[start:end]
                return json.loads(json_str)
            return None
        except:
            return None

    @staticmethod
    def aggregate_field(data: List[Dict], field: str, operation: str) -> Dict:
        """Perform aggregation on a field."""
        values = []
        for item in data:
            if isinstance(item, dict):
                # Extract nested fields using dot notation
                value = item
                for key in field.split('.'):
                    if isinstance(value, dict) and key in value:
                        value = value[key]
                    else:
                        value = None
                        break
                if value is not None and isinstance(value, (int, float)):
                    values.append(value)

        if not values:
            return {"error": f"No numeric values found for field {field}"}

        result = {
            "field": field,
            "count": len(values)
        }

        if operation == 'max':
            result["max"] = max(values)
        elif operation == 'min':
            result["min"] = min(values)
        elif operation == 'avg':
            result["average"] = mean(values)
        elif operation == 'sum':
            result["sum"] = sum(values)
        elif operation == 'all':
            result.update({
                "max": max(values),
                "min": min(values),
                "average": mean(values),
                "sum": sum(values)
            })

        return result

    @staticmethod
    def group_by(data: List[Dict], group_field: str, agg_field: str, operation: str) -> Dict:
        """Group data by field and perform aggregation."""
        groups = defaultdict(list)
        
        for item in data:
            if isinstance(item, dict):
                # Get group value
                group_value = item
                for key in group_field.split('.'):
                    if isinstance(group_value, dict) and key in group_value:
                        group_value = group_value[key]
                    else:
                        group_value = None
                        break
                
                # Get aggregation value
                agg_value = item
                for key in agg_field.split('.'):
                    if isinstance(agg_value, dict) and key in agg_value:
                        agg_value = agg_value[key]
                    else:
                        agg_value = None
                        break
                
                if group_value is not None and agg_value is not None:
                    groups[str(group_value)].append(agg_value)
        
        result = {}
        for group, values in groups.items():
            if operation == 'max':
                result[group] = max(values)
            elif operation == 'min':
                result[group] = min(values)
            elif operation == 'avg':
                result[group] = mean(values)
            elif operation == 'sum':
                result[group] = sum(values)
            elif operation == 'count':
                result[group] = len(values)
        
        return result

class PerformanceMonitor:
    def __init__(self):
        self.metrics = defaultdict(list)
    
    def record_metric(self, name: str, value: float):
        self.metrics[name].append(value)
    
    def get_average(self, name: str) -> float:
        values = self.metrics[name]
        return sum(values) / len(values) if values else 0
    
    def reset(self):
        self.metrics.clear()

class IndexingHelper:
    def __init__(self, redis_url: str):
        self.redis_client = redis.from_url(redis_url)
        self._vector_store = None
        self.performance = PerformanceMonitor()
        
        # Initialize specialized readers for each file type
        self.readers = {
            'txt': FlatReader(),
            'docx': DocxReader(),
            'pdf': PDFReader(),
            'json': FlatReader(),
            'csv': PandasCSVReader(
                concat_rows=False,  # Create separate document for each row
                col_joiner=", ",
                pandas_config={'encoding': 'utf-8'}
            ),
            'xlsx': PandasExcelReader(
                concat_rows=False,
                sheet_name=None  # Process all sheets
            ),
            # 'pptx': PptxReader()
        }
    
    def initialize_vector_store(self, weaviate_client: Any):
        """Initialize the vector store."""
        self._vector_store = WeaviateVectorStore(
            weaviate_client=weaviate_client,
            index_name="Document",
            text_key="content"
        )
        print("Vector store initialized:", self._vector_store)

    def process_document(self, file_path: str, weaviate_client: Any, embedding_model: OllamaEmbeddings, chunk_size: Optional[int] = None) -> bool:
        """Optimized document processing with specialized readers."""
        try:
            # Initialize the vector store
            self.initialize_vector_store(weaviate_client)
            
            # Get appropriate reader based on file extension
            file_extension = Path(file_path).suffix.lower()[1:]
            reader = self.readers.get(file_extension)
            
            if not reader:
                raise ValueError(f"Unsupported file type: {file_extension}")
            
            # Load document with enhanced metadata
            documents = reader.load_data(
                Path(file_path),
                extra_info={
                    'source': Path(file_path).name,
                    'file_type': file_extension,
                    'timestamp': datetime.datetime.now(datetime.timezone.utc).isoformat(),  # Use timezone-aware datetime
                    'processor': reader.__class__.__name__,
                    'page_count': self._get_page_count(file_path) if file_extension == 'pdf' else None
                }
            )
            
            print(f"Loaded {len(documents)} documents from {file_path}")

            # Process each document individually
            for chunk_index, document in enumerate(documents):
                print(f"Processing document: {document}")  # Log the entire document object
                print(f"Document attributes: {vars(document)}")  # Log the attributes of the document
                print("Using model:", OLLAMA_MODEL)
                # Access the correct method for generating embeddings
                embedding = embedding_model.embed_documents(document.text)  # Change this if the method is different
                print("embedding->>>>>", embedding)
                print("embedding->>>>>" , len(embedding), len(embedding[0]))

                # Store each document chunk in Weaviate directly
                # weaviate_client.data.create(
                #     data={
                #         "content": document.content,
                #         "source": document.source,
                #         "chunk_index": chunk_index,
                #         "status": "processed"
                #     },
                #     class_name="Document",
                #     vector=embedding
                # )

                    # Create a collection if it doesn't exist
                if not weaviate_client.collections.exists("Document"):
                    print("document collection do not exist creating one ------------")
                    weaviate_client.collections.create(
                        name="Document",
                        properties=[
                            {"name": "content", "dataType": "text"},
                            {"name": "source", "dataType": "text"},
                            {"name": "chunk_index", "dataType": "int"},
                            {"name": "status", "dataType": "text"}
                        ]
                    )

                # In your process_document method, replace the weaviate_client.data.create() call with:
                document_collection = weaviate_client.collections.get("Document")
                document_collection.data.insert(
                    properties={
                        "content": document.text,
                        "source": document.metadata.get('source'),
                        "chunk_index": chunk_index,
                        "status": "processed"
                    },
                    vector=embedding[0]
                )

            
            print(f"Successfully processed and indexed {len(documents)} documents.")
            return True
            
        except Exception as e:
            print(f"Error processing document: {str(e)}")
            return False

    def _get_page_count(self, file_path: str) -> Optional[int]:
        """Get page count for PDF files."""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PdfReader(file)
                return len(pdf_reader.pages)
        except:
            return None

    def query_document(self, query: str, document_name: str, limit: int = 3, unique: bool = False) -> List[Dict]:
        """Query documents from the vector store.
        
        Args:
            query: The search query
            document_name: Name of the source document to filter by
            limit: Maximum number of results to return
            unique: If True, return unique results only (currently not implemented)
        
        Returns:
            List of dictionaries containing the search results
        """
        try:
            if not self._vector_store:
                logger.error("Vector store not initialized")
                return []

            # Use similarity_search with the query
            nodes = self._vector_store.similarity_search(
                query,
                k=limit,
                where_filter={
                    "path": ["source"],
                    "operator": "Equal",
                    "valueString": document_name
                }
            )
            
            # Format results to match our document structure
            results = []
            for node in nodes:
                result = {
                    "content": node.text,
                    "source": node.metadata.get('source'),
                    "chunk_index": node.metadata.get('chunk_index'),
                    "status": node.metadata.get('status', 'retrieved'),
                    "score": node.score if hasattr(node, 'score') else None
                }
                results.append(result)
            
            logger.info(f"Found {len(results)} results for query in document {document_name}")
            return results
            
        except Exception as e:
            logger.error(f"Error querying document: {str(e)}")
            return []

    

    def query_json_document(self, 
                          query: Dict[str, Any], 
                          document_name: str, 
                          limit: int = 3) -> Dict:
        """Query JSON document with aggregations."""
        try:
            # Get base results
            results = self.query_document(
                query.get('text', ''),
                document_name,
                limit
            )
            
            # Extract JSON data from results
            json_data = []
            for result in results:
                json_obj = JSONQueryHelper.extract_json_fields(result['content'])
                if json_obj:
                    json_data.append(json_obj)
            
            response = {
                'matches': results,
                'json_analysis': {}
            }
            
            # Handle aggregations
            if 'aggregate' in query:
                agg_config = query['aggregate']
                field = agg_config.get('field')
                operation = agg_config.get('operation', 'all')
                
                if field:
                    response['json_analysis']['aggregation'] = \
                        JSONQueryHelper.aggregate_field(json_data, field, operation)
            
            # Handle grouping
            if 'group_by' in query:
                group_config = query['group_by']
                group_field = group_config.get('field')
                agg_field = group_config.get('aggregate_field')
                operation = group_config.get('operation', 'count')
                
                if group_field and agg_field:
                    response['json_analysis']['groups'] = \
                        JSONQueryHelper.group_by(json_data, group_field, agg_field, operation)
            
            return response
            
        except Exception as e:
            logger.error(f"Error querying JSON document: {str(e)}")
            return {'error': str(e)} 

    def close(self):
        if self._vector_store:
            self._vector_store.client.close()  # Close the Weaviate client connection
            print("Weaviate client connection closed.")

# Example usage of WeaviateHelper
if __name__ == "__main__":
    client = WeaviateHelper.get_client()
    if client:
        # Perform operations with the client
        pass 